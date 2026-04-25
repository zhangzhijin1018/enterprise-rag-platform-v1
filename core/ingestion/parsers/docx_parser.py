"""DOCX 解析器模块。

Word 文档在企业知识库里很常见，但它的问题也很典型：

- 标题、正文、列表、表格混在同一个文件里
- 视觉上能看懂的层级，程序如果不显式保留，后续就会被“拍扁”
- 纯粹抽取连续文本会丢失“这段内容属于哪个章节”的语义

因此这里不是简单把所有段落直接拼接，而是把 DOCX 里的结构线索转成更适合 RAG 的文本：

1. 标题样式转成 Markdown 标题，便于后续 chunker 识别层级
2. 正文和列表项显式带上 `[section:章节名]`，避免脱离标题后语义丢失
3. 表格转成“表名 + 行标题 + 字段:值”的一维文本，方便 embedding 和 keyword 检索
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document as DocxDocument

from core.ingestion.cleaners.text_cleaner import clean_text
from core.models.document import Document

from .base import BaseParser


_HEADING_LEVEL_RE = re.compile(r"Heading\s*(\d+)", re.IGNORECASE)


def _heading_level_from_style(style_name: str) -> int | None:
    """从 Word 段落样式里提取标题层级。

    `python-docx` 暴露的段落样式通常是 `Heading 1`、`Heading 2` 这类名字。
    这里把它们映射到 Markdown 风格的 `#` 层级，是为了统一不同文件格式在下游的结构表达。
    """

    match = _HEADING_LEVEL_RE.search(style_name)
    if not match:
        return None
    return max(1, min(6, int(match.group(1))))


def _is_list_style(style_name: str) -> bool:
    """判断当前段落是否更像列表项。

    列表在制度、SOP、巡检步骤文档中非常常见。
    把它和普通正文区分开，后续更容易保留“步骤/枚举项”的阅读语气。
    """

    lowered = style_name.strip().lower()
    return "list" in lowered or "bullet" in lowered or "number" in lowered


class DocxParser(BaseParser):
    """把 `.docx` 文件解析成带结构标记的统一 `Document`。"""

    def parse(self, path: Path, source: str) -> Document:
        """解析 DOCX 并保留章节、列表、表格结构。

        这一步的目标不是复原 Word 的完整版式，而是提取“对检索最有价值的结构信息”：

        - 标题层级
        - 章节归属
        - 列表项语义
        - 表格字段关系
        """

        d = DocxDocument(str(path))
        lines: list[str] = []
        # 记录最近一次遇到的标题。
        # 后续正文即使脱离标题本身，也能通过 `[section:xxx]` 继续保留章节语义。
        current_section = ""
        for p in d.paragraphs:
            text = (p.text or "").strip()
            # 空段落通常只是版式占位，对检索没有帮助，直接跳过避免噪声。
            if not text:
                continue
            style = (p.style.name if p.style else "") or ""
            level = _heading_level_from_style(style)
            if level is not None:
                # 标题本身既要成为文本内容，也要更新后续正文的 section 上下文。
                current_section = text
                lines.append(f"{'#' * level} {text}")
            elif _is_list_style(style):
                # 列表项常常对应步骤、要点、限制条件。
                # 如果能识别出所属章节，就显式带上 section，方便后续召回时理解上下文。
                if current_section:
                    lines.append(f"[section:{current_section}] - {text}")
                else:
                    lines.append(f"- {text}")
            else:
                # 普通正文同样尽量补 section 标签，减少“正文离标题太远后语义变弱”的问题。
                if current_section:
                    lines.append(f"[section:{current_section}] {text}")
                else:
                    lines.append(text)

        # `python-docx` 的表格不在 `paragraphs` 里，需要单独读取。
        # 这里把二维表格转成更适合 RAG 的一维字段文本，核心目标是保留“列名 -> 值”的关系。
        for table_idx, table in enumerate(d.tables, start=1):
            rows = table.rows
            if not rows:
                continue
            # 默认把第一行视作表头；若单元格为空，则补一个稳定占位名，避免后续丢字段。
            headers = [clean_text(cell.text).strip() or f"column_{idx + 1}" for idx, cell in enumerate(rows[0].cells)]
            lines.append(f"## Table {table_idx}")
            for row_idx, row in enumerate(rows[1:], start=1):
                values = [clean_text(cell.text).strip() for cell in row.cells]
                # 整行都为空时通常没有信息量，直接跳过。
                if not any(values):
                    continue
                lines.append(f"### Table {table_idx} Row {row_idx}")
                for col_idx, value in enumerate(values):
                    if not value:
                        continue
                    # 列数不齐时兜底生成列名，保证文本表达稳定，不因为脏数据报错。
                    header = headers[col_idx] if col_idx < len(headers) else f"column_{col_idx + 1}"
                    lines.append(f"{header}: {value}")

        # 最后统一再清洗一次，确保段落和表格拼接后没有多余空白与噪声字符。
        content = clean_text("\n".join(lines))
        return Document(
            # `doc_id` 由更上游的 ingestion 流程生成；parser 只负责“把内容读出来”。
            doc_id="",
            source=source,
            # 对本地文件来说，文件名通常是最稳定的默认标题兜底。
            title=path.stem,
            content=content,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            # DOCX parser 当前不额外补充文档级 metadata，后续可由 metadata extractor 继续增强。
            metadata={},
        )
